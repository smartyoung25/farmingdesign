"""
문서청킹 공통 라이브러리 — `문서청킹구조_설계공사비시방서.md` 스키마의 태깅·청킹 로직.
`build_document_chunks.py`(5개 사례 파일럿)와 `build_document_chunks_full.py`(전체 확장)가
같은 로직을 공유하도록 분리했다(2026-07-24, 파일럿 검증 후 전면 확장하며 리팩터링).

주의: 이 인덱스는 엔진의 계산 출처가 아니다. 엔진 상수 승격은 기존 실측 확인
절차(레지스트리 등록 + 드리프트 가드 테스트)를 그대로 거친다.
"""
import json
import os
import re

import pdfplumber
import openpyxl

ROOT = os.path.dirname(os.path.abspath(__file__))
SPEC_DIR = os.path.join(ROOT, "스마트팜스펙")
FACILITY_DIR = os.path.join(ROOT, "시설평가")

TOPIC_KEYWORDS = {
    "부지": r"부지|지반|진입로|구획|성토|절토|경사|필지|배수로|배수계획|배수시설|용배수",
    "기후": r"적설|풍속|기후|강설|동상|결빙|내재해",
    "작물특성": r"작물|작목|토마토|딸기|무화과|재배작물|정식|파종",
    "냉난방": r"난방|냉방|보일러|열관류율|보온|커튼|열원|온풍기|냉방기|열교환",
    "광량": r"광투과|차광|보광|일사|광량|투광",
    "CO2": r"CO2|이산화탄소|탄산가스",
    "생육관리": r"생육|병해충|방제|양액|관비|정지|유인|착과",
    "운영": r"ICT|제어기|자동화|인건비|운영비|전기요금|에너지비용",
    "유지보수": r"하자|보수|점검|유지관리|내구연한|보증기간",
    "구조/자재": r"철골|골조|기초|파이프|트러스|부재|피복재|비닐|유리|파형강판|하중|설계기준|KDS",
    "비용/공정": r"공사비|내역서|단가|수량|재료비|노무비|경비|공정표|공기|일정|합계",
}
TOPIC_PATTERNS = {k: re.compile(v) for k, v in TOPIC_KEYWORDS.items()}

CAPEX_KEYWORD = re.compile(r"공사비|내역서|재료비|노무비|경비|원가계산")
STRUCT_KEYWORD = re.compile(r"적설|풍속|하중|구조계산|KDS")

# 구조계산서 등 PDF에서 폰트/인증서 메타데이터가 텍스트로 새어나오는 노이즈 패턴
# (2026-07-24 파일럿에서 발견: "Certified by Gen MODEL DATA PROFILE..." 반복)
NOISE_PATTERN = re.compile(r"Certifi|MODEL DATA|GDLicense|Robot\s*Structural|autodesk", re.I)

# 공사비내역서/설계예산서 표에서 "공종 헤더 행"으로 볼 패턴 (예: "1-1. 기초공사", "0102. 기초공사")
SECTION_HEADER_PATTERN = re.compile(r"^\d+([.\-]\d+)*\.?\s*[가-힣]")


def looks_like_noise(text):
    return len(NOISE_PATTERN.findall(text)) >= 2


def is_header_row(cells):
    if not cells or len(cells) > 2:
        return False
    first = cells[0].strip()
    return bool(SECTION_HEADER_PATTERN.match(first))


def tag_topics(text):
    return [name for name, pat in TOPIC_PATTERNS.items() if pat.search(text)]


def engine_link_for(doc_type, text):
    links = []
    if doc_type in ("공사비내역서", "설계예산서", "수량산출서") and CAPEX_KEYWORD.search(text):
        links.append("CAPEX_MAJOR_CATEGORIES")
    if doc_type == "구조계산서" and STRUCT_KEYWORD.search(text):
        links += ["SPEC_TABLE", "REGION_DESIGN_LOAD"]
    return links or None


def clean(t):
    return re.sub(r"\s+", " ", t or "").strip()


class ChunkWriter:
    def __init__(self):
        self.chunks = []
        self.skipped_files = []  # (path, reason)

    def add(self, case_name, doc_type, doc_subtype, source_file, page_or_section,
            text, evidence_status="실측", extraction_quality="normal", duplicate_of=None,
            section_context=None):
        text = clean(text)
        noise = looks_like_noise(text)
        tagging_text = text if not noise else ""
        if section_context:
            tagging_text = clean(section_context) + " " + tagging_text
        chunk_id = f"{case_name}_{doc_type}_{doc_subtype or ''}_{page_or_section}".replace(" ", "")
        rec = {
            "chunk_id": chunk_id,
            "source_file": os.path.relpath(source_file, ROOT).replace("\\", "/"),
            "case_name": case_name,
            "doc_type": doc_type,
            "doc_subtype": doc_subtype,
            "page_or_section": page_or_section,
            "section_context": clean(section_context) if section_context else None,
            "topic_tags": tag_topics(tagging_text),
            "content_summary": text[:200],
            "text_len": len(text),
            "evidence_status": evidence_status,
            "extraction_quality": "noise_suspected" if noise else extraction_quality,
            "engine_link": engine_link_for(doc_type, tagging_text),
            "duplicate_of": duplicate_of,
        }
        self.chunks.append(rec)

    def skip(self, path, reason):
        self.skipped_files.append((os.path.relpath(path, ROOT).replace("\\", "/"), reason))


def chunk_pdf_by_page(w, path, case_name, doc_type, doc_subtype, evidence_status="실측",
                       quality_note=None, duplicate_of=None):
    with pdfplumber.open(path) as pdf:
        n_pages = len(pdf.pages)
        n_blank = 0
        for i, page in enumerate(pdf.pages):
            t = page.extract_text() or ""
            if not t.strip():
                n_blank += 1
                continue
            w.add(case_name, doc_type, doc_subtype, path, f"p{i+1}", t,
                  evidence_status=evidence_status,
                  extraction_quality=quality_note or "normal",
                  duplicate_of=duplicate_of)
    return n_pages, n_blank


def chunk_pdf_table_or_lines(w, path, case_name, doc_type, doc_subtype, evidence_status="실측",
                              duplicate_of=None):
    n_pages_with_table = 0
    n_pages_line_fallback = 0
    current_section = None
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            if tables:
                n_pages_with_table += 1
                for ti, table in enumerate(tables):
                    for ri, row in enumerate(table):
                        cells = [clean(c) for c in row if c]
                        if not cells:
                            continue
                        row_text = " | ".join(cells)
                        if len(row_text) < 2:
                            continue
                        if is_header_row(cells):
                            current_section = row_text
                            ctx = None
                        else:
                            ctx = current_section
                        w.add(case_name, doc_type, doc_subtype, path, f"p{i+1}-t{ti+1}-r{ri+1}",
                              row_text, evidence_status=evidence_status,
                              extraction_quality="table_extract", duplicate_of=duplicate_of,
                              section_context=ctx)
            else:
                t = page.extract_text() or ""
                if not t.strip():
                    continue
                n_pages_line_fallback += 1
                for li, line in enumerate(t.split("\n")):
                    if len(line.strip()) < 2:
                        continue
                    if is_header_row([line.strip()]):
                        current_section = line.strip()
                        ctx = None
                    else:
                        ctx = current_section
                    w.add(case_name, doc_type, doc_subtype, path, f"p{i+1}-line{li+1}",
                          line, evidence_status=evidence_status,
                          extraction_quality="line_fallback", duplicate_of=duplicate_of,
                          section_context=ctx)
    return n_pages_with_table, n_pages_line_fallback


def chunk_xlsx(w, path, case_name, doc_type, doc_subtype, evidence_status="실측"):
    wb = openpyxl.load_workbook(path, data_only=True)
    n_rows = 0
    for ws in wb.worksheets:
        current_section = None
        for ri, row in enumerate(ws.iter_rows(values_only=True), start=1):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if not cells:
                continue
            row_text = " | ".join(cells)
            if len(row_text) < 2:
                continue
            n_rows += 1
            if is_header_row(cells):
                current_section = row_text
                ctx = None
            else:
                ctx = current_section
            w.add(case_name, doc_type, doc_subtype, path, f"{ws.title}-r{ri}",
                  row_text, evidence_status=evidence_status, extraction_quality="xlsx_row",
                  section_context=ctx)
    return n_rows


def chunk_xls_legacy(w, path, case_name, doc_type, doc_subtype, evidence_status="실측"):
    import pandas as pd
    n_rows = 0
    try:
        sheets = pd.read_excel(path, sheet_name=None, header=None, engine="xlrd")
    except Exception as e:
        w.skip(path, f"xls 읽기 실패: {e}")
        return 0
    for sheet_name, df in sheets.items():
        current_section = None
        for ri, row in df.iterrows():
            cells = [str(c) for c in row.tolist() if str(c) not in ("nan", "None", "")]
            if not cells:
                continue
            row_text = " | ".join(cells)
            if len(row_text) < 2:
                continue
            n_rows += 1
            if is_header_row(cells):
                current_section = row_text
                ctx = None
            else:
                ctx = current_section
            w.add(case_name, doc_type, doc_subtype, path, f"{sheet_name}-r{ri+1}",
                  row_text, evidence_status=evidence_status, extraction_quality="xls_row",
                  section_context=ctx)
    return n_rows


def chunk_docx(w, path, case_name, doc_type, doc_subtype, evidence_status="실측"):
    import docx
    d = docx.Document(path)
    n = 0
    current_section = None
    HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Title"}
    for pi, para in enumerate(d.paragraphs):
        t = clean(para.text)
        if len(t) < 2:
            continue
        is_heading = para.style.name in HEADING_STYLES or bool(SECTION_HEADER_PATTERN.match(t))
        if is_heading:
            current_section = t
            ctx = None
        else:
            ctx = current_section
        n += 1
        w.add(case_name, doc_type, doc_subtype, path, f"para{pi+1}", t,
              evidence_status=evidence_status, extraction_quality="docx_para",
              section_context=ctx)
    for ti, table in enumerate(d.tables):
        for ri, row in enumerate(table.rows):
            cells = [clean(c.text) for c in row.cells if clean(c.text)]
            if not cells:
                continue
            row_text = " | ".join(cells)
            n += 1
            w.add(case_name, doc_type, doc_subtype, path, f"t{ti+1}-r{ri+1}",
                  row_text, evidence_status=evidence_status, extraction_quality="docx_table_row")
    return n


def write_outputs(w, jsonl_path, summary_path, extra_log=None):
    with open(jsonl_path, "w", encoding="utf-8") as f:
        for c in w.chunks:
            f.write(json.dumps(c, ensure_ascii=False) + "\n")

    from collections import Counter
    by_doc_type = Counter(c["doc_type"] for c in w.chunks)
    by_case = Counter(c["case_name"] for c in w.chunks)
    by_quality = Counter(c["extraction_quality"] for c in w.chunks)
    topic_counter = Counter()
    for c in w.chunks:
        for t in c["topic_tags"]:
            topic_counter[t] += 1
    untagged = sum(1 for c in w.chunks if not c["topic_tags"])
    linked = sum(1 for c in w.chunks if c["engine_link"])

    with open(summary_path, "w", encoding="utf-8") as f:
        if extra_log:
            f.write("=== 처리 로그 ===\n")
            for line in extra_log:
                f.write(line + "\n")
            f.write("\n")

        f.write(f"=== 총 청크 수: {len(w.chunks)} ===\n")
        f.write(f"=== 처리 실패/건너뛴 파일 수: {len(w.skipped_files)} ===\n")

        f.write("\n-- doc_type별 청크 수 --\n")
        for k, v in by_doc_type.most_common():
            f.write(f"  {k}: {v}\n")
        f.write("\n-- case별 청크 수 (상위 30) --\n")
        for k, v in by_case.most_common(30):
            f.write(f"  {k}: {v}\n")
        f.write("\n-- extraction_quality별 청크 수 --\n")
        for k, v in by_quality.most_common():
            f.write(f"  {k}: {v}\n")
        f.write("\n-- 주제축별 태깅 건수 (청크 1개가 복수 축에 걸릴 수 있음) --\n")
        for k, v in topic_counter.most_common():
            f.write(f"  {k}: {v}\n")
        f.write(f"\n미분류(태그 0개) 청크 수: {untagged} / {len(w.chunks)}"
                f" ({untagged/len(w.chunks)*100:.1f}%)\n" if w.chunks else "\n청크 없음\n")
        f.write(f"\nengine_link 있음: {linked} / {len(w.chunks)}\n")

        if w.skipped_files:
            f.write("\n-- 건너뛴 파일 (근거 없이 채우지 않고 명시적으로 제외) --\n")
            for path, reason in w.skipped_files:
                f.write(f"  {path}: {reason}\n")

    print(f"완료: 청크 {len(w.chunks)}개 -> {jsonl_path}")
    print(f"건너뛴 파일 {len(w.skipped_files)}개")
    print(f"요약 -> {summary_path}")
