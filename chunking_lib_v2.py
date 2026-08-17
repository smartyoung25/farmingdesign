"""
문서청킹 공통 라이브러리 — `문서청킹구조_설계공사비시방서.md` 스키마의 태깅·청킹 로직.
`build_document_chunks.py`(5개 사례 파일럿)와 `build_document_chunks_full.py`(전체 확장)가
같은 로직을 공유하도록 분리했다(2026-07-24, 파일럿 검증 후 전면 확장하며 리팩터링).

주의: 이 인덱스는 엔진의 계산 출처가 아니다. 엔진 상수 승격은 기존 실측 확인
절차(레지스트리 등록 + 드리프트 가드 테스트)를 그대로 거친다.
"""
import json
import os
import re

import pdfplumber
import openpyxl

ROOT = os.path.dirname(os.path.abspath(__file__))
SPEC_DIR = os.path.join(ROOT, "스마트팜스펙")
FACILITY_DIR = os.path.join(ROOT, "시설평가")


# ── P2-23(2026-08-17): 파이프라인 의존성·산출물 급감 가드 ──────────────
# 배경: 2026-08-16 xlrd 유실 사건(작업지시서 10-4절) — 의존성 부재가 skip 사유로만
# 기록된 채 실행이 계속돼 .xls 파일들이 조용히 0청크가 됐고(기존 대비 2,345청크
# 증발), 병합 후 총량을 이전 실행과 수동 대조하고서야 발견됐다. 2026-08-17엔
# pdfplumber도 같은 방식으로 환경에서 유실돼 있었음이 재확인됨. 조용한 열화를
# 실행 전(의존성 확인)·병합 시(총량 대조) 두 지점의 명시적 중단으로 바꾼다.
PIPELINE_DEPENDENCIES = ("pdfplumber", "pypdf", "openpyxl", "xlrd", "pandas", "docx")


def assert_pipeline_dependencies(deps=PIPELINE_DEPENDENCIES):
    """필수 파서 패키지가 하나라도 없으면 RuntimeError로 즉시 중단한다.
    없는 채 진행하면 해당 포맷 파일 전부가 '읽기 실패' skip으로 조용히 사라진다."""
    import importlib
    missing = []
    for name in deps:
        try:
            importlib.import_module(name)
        except ImportError:
            missing.append(name)
    if missing:
        raise RuntimeError(
            "청킹 파이프라인 필수 패키지 유실: " + ", ".join(missing)
            + " — pip install 후 재실행할 것 (P2-23 가드: 의존성 없이 진행하면 "
              "해당 포맷이 조용히 0청크가 된다. 10-4절 xlrd 사건 참고)")
    return True


def chunk_count_regression_guard(new_count, index_path, allow_shrink=None):
    """병합 결과 총 청크 수가 기존 인덱스보다 줄면 RuntimeError로 중단한다.
    의도적 축소(파일 삭제·스코프 변경)는 환경변수 CHUNK_ALLOW_SHRINK=1 로만 통과."""
    if allow_shrink is None:
        allow_shrink = os.environ.get("CHUNK_ALLOW_SHRINK") == "1"
    if not os.path.exists(index_path):
        return True  # 첫 실행 — 비교 기준 없음
    with open(index_path, encoding="utf-8") as f:
        old_count = sum(1 for line in f if line.strip())
    if new_count < old_count and not allow_shrink:
        raise RuntimeError(
            f"병합 결과 청크 수 급감: 기존 {old_count} → 신규 {new_count}"
            f"({old_count - new_count} 감소) — 의존성 유실/부분 처리 의심"
            "(10-4절 사건과 동일 패턴). 의도적 축소가 맞으면 "
            "CHUNK_ALLOW_SHRINK=1 환경변수로 재실행할 것")
    return True


# ── P2-14(2026-08-17): 매니페스트(4상태 diff) · 수동보정 오버레이 ──────────
# 방법론 문서 6-1·6-2·6-4절 구현. 매니페스트는 파일 단위 처리상태 대장으로,
# 재실행 시 지문(sha256+mtime/size) 대조로 무변경/신규/변경/삭제를 판정한다
# (종전엔 파트파일 존재 여부만 봐서 원본이 바뀌거나 지워져도 감지 불가).
# 오버레이는 사람이 고친 태그를 원본 인덱스 밖(청킹_오버레이.jsonl)에 쌓아
# 재실행에도 유실되지 않게 한다 — 최종 뷰 = 자동 인덱스 ⊕ 오버레이(오버레이 우선).
import hashlib as _hashlib

MANIFEST_VERSION = "1.0"
MANIFEST_PATH = os.path.join(ROOT, "청킹_매니페스트.json")
OVERLAY_PATH = os.path.join(ROOT, "청킹_오버레이.jsonl")
# 오버레이가 덮어쓸 수 있는 필드 — 식별자(chunk_id·source_file)와 지문성 필드는 불가
OVERLAY_ALLOWED_FIELDS = {"topic_tags", "consulting_tags", "doc_type", "doc_subtype",
                          "case_name", "engine_link", "content_summary", "evidence_status"}


def file_fingerprint(path):
    """파일 지문 — sha256(내용) + mtime_ns/size(보조)."""
    h = _hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(1 << 20), b""):
            h.update(block)
    st = os.stat(path)
    return {"sha256": h.hexdigest(), "mtime": st.st_mtime_ns, "size": st.st_size}


def load_manifest(path=MANIFEST_PATH):
    if os.path.exists(path):
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    return {"manifest_version": MANIFEST_VERSION, "last_run": None, "files": {}}


def save_manifest(manifest, path=MANIFEST_PATH):
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=1)
    os.replace(tmp, path)  # 원자적 저장(파트파일과 동일한 체크포인트 규율)


def classify_file_states(current, manifest_files):
    """4상태 diff(방법론 6-2절). current: {relpath: fingerprint dict},
    manifest_files: 매니페스트의 files 딕셔너리. 반환: 상태별 relpath 리스트."""
    states = {"new": [], "changed": [], "unchanged": [], "deleted": []}
    for rel, fp in current.items():
        ent = manifest_files.get(rel)
        if ent is None:
            states["new"].append(rel)
        elif ent.get("sha256") != fp["sha256"]:
            states["changed"].append(rel)
        else:
            states["unchanged"].append(rel)
    for rel, ent in manifest_files.items():
        if rel not in current and ent.get("status") != "tombstoned":
            states["deleted"].append(rel)
    return states


def load_overlay(path=OVERLAY_PATH):
    """오버레이 항목 목록. 파일이 없거나 비어 있으면 빈 목록(보정이 아직 없다는 뜻)."""
    entries = []
    if not os.path.exists(path):
        return entries
    with open(path, encoding="utf-8") as f:
        for ln, line in enumerate(f, 1):
            line = line.strip()
            if not line:
                continue
            try:
                entries.append(json.loads(line))
            except Exception as e:
                raise RuntimeError(f"오버레이 {ln}행 JSON 파싱 실패: {e} — 손보정 파일이라 자동 무시하지 않고 중단")
    return entries


def apply_overlay(chunks, overlay_entries):
    """자동 인덱스 청크에 오버레이를 적용한 '최종 뷰'를 만든다(원본 리스트 불변).
    매칭 키는 chunk_id(+ 오버레이에 source_file이 있으면 함께 대조 — chunk_id는
    파일 간 충돌이 가능하므로 모호하면 source_file로 한정할 것). 적용된 청크엔
    manual_overlay=True 를 표시한다. 반환: (뷰 리스트, 적용 건수, 미매칭 오버레이 수)."""
    by_id = {}
    for e in overlay_entries:
        by_id.setdefault(e["chunk_id"], []).append(e)
    out, n_applied, matched = [], 0, set()
    for c in chunks:
        hits = [e for e in by_id.get(c["chunk_id"], [])
                if not e.get("source_file") or e["source_file"] == c.get("source_file")]
        if not hits:
            out.append(c)
            continue
        merged = dict(c)
        for e in hits:
            for k, v in (e.get("fields") or {}).items():
                if k in OVERLAY_ALLOWED_FIELDS:
                    merged[k] = v
            matched.add(id(e))
        merged["manual_overlay"] = True
        out.append(merged)
        n_applied += 1
    n_unmatched = sum(1 for e in overlay_entries if id(e) not in matched)
    return out, n_applied, n_unmatched

# ── 주제축(도메인) 태깅 규칙 ──
# 2026-07-24 개정: 사용자 지정 9개 공학 도메인 축으로 재구성한다.
#   부지 · 토목 · 시설 · 장비 · 전기 · 통신 · 재배환경 · 운영 · 사후관리
# 시방서/설계서/견적서/도면 텍스트(행·페이지)를 이 9개 축으로 분해 태깅한다.
# 한 청크가 복수 축에 걸릴 수 있다(예: "난방기"는 장비이자 재배환경).
# (이전 파일럿의 축: 부지/기후/작물특성/냉난방/광량/CO2/생육관리/운영/유지보수/
#  구조자재/비용공정 — 개정판에서 재배환경·시설·장비 등으로 통합·재배치됨)
TOPIC_KEYWORDS = {
    # 대지·필지·진입·측량 등 '땅' 관련
    "부지": r"부지|필지|대지|지목|용지|진입로|경계|구획|배치도|위치도|측량|녹지|부지경계|지적",
    # 지반·토공·배수·기초토목 등 '토목공사'
    "토목": r"토목|지반|성토|절토|터파기|되메우기|다짐|굴착|잡석|버림\s*콘크리트|기초\s*콘크리트|"
            r"배수|우수|맨홀|측구|집수정|옹벽|포장|부지조성|정지작업|경사|사면|암거|토사|되메움",
    # 온실 구조체·골조·피복 등 '시설(구조물)'
    "시설": r"온실|하우스|골조|철골|트러스|서까래|기둥|중방|들보|파이프|피복|유리|비닐|폴리|외피|"
            r"지붕|천창|측창|측벽|박공|연동|단동|구조물|파형강판|스크린|보온\s*커튼|차광막|개폐부|용마루",
    # 설비·기자재·냉난방장치 등 '장비'
    "장비": r"기자재|장비|설비|양액기|양액|관수|점적|분수|환기팬|유동팬|순환팬|배기팬|난방기|보일러|"
            r"히트\s*펌프|온풍기|냉방기|칠러|제어기|컨트롤러|센서|개폐기|감속기|모터|FCU|팬코일|"
            r"CO2\s*발생|탄산가스\s*발생|약액|배지|양액\s*탱크|히트펌프",
    # 수배전·배선·접지 등 '전기'
    "전기": r"전기|수배전|배전|분전반|배선|접지|차단기|전력|인입|케이블|전등|조명|동력|콘센트|전압|"
            r"한전|누전|간선|전기공사|전기설비|수전|변압기|전기배관",
    # 네트워크·통신 인프라 '통신'
    "통신": r"통신|네트워크|인터넷|유선|무선|LAN|랜|광케이블|광\s*통신|IoT|사물인터넷|원격|공유기|"
            r"라우터|데이터\s*통신|CCTV|IP\s*카메라|통신배관|통신설비",
    # 작물·환경(온습도·광·CO2·양액·냉난방 운용) '재배환경'
    "재배환경": r"재배|작물|작목|토마토|딸기|파프리카|무화과|생육|정식|착과|유인|파종|"
                r"온도|습도|일사|광량|투광|차광|보광|CO2|이산화탄소|양액\s*농도|관비|병해충|방제|"
                r"VPD|보온|냉방|난방|환기|결로|과습|열관류율|열원|생육관리",
    # 운영·인건·에너지비 등 '운영'
    # 농장 '운영' 단계 비용(건설 노무비/재료비는 제외 — 그건 doc_type=견적서/내역서로 포착)
    "운영": r"운영비|운영관리|가동비|가동률|인건비|관리비|전기요금|연료비|에너지비용|난방비|"
            r"유지비|경영|자동화|ICT|복합환경제어",
    # 하자·보수·점검·보증 등 '사후관리'
    "사후관리": r"하자|보수|사후관리|유지관리|점검|보증|보증기간|A\s*/?\s*S|AS\b|내구연한|소모품|"
                r"교체|정기점검|워런티|유지보수",
}
TOPIC_PATTERNS = {k: re.compile(v) for k, v in TOPIC_KEYWORDS.items()}

CAPEX_KEYWORD = re.compile(r"공사비|내역서|재료비|노무비|경비|원가계산")
STRUCT_KEYWORD = re.compile(r"적설|풍속|하중|구조계산|KDS")

# 구조계산서 등 PDF에서 폰트/인증서 메타데이터가 텍스트로 새어나오는 노이즈 패턴
# (2026-07-24 파일럿에서 발견: "Certified by Gen MODEL DATA PROFILE..." 반복)
NOISE_PATTERN = re.compile(r"Certifi|MODEL DATA|GDLicense|Robot\s*Structural|autodesk", re.I)

# 공사비내역서/설계예산서 표에서 "공종 헤더 행"으로 볼 패턴 (예: "1-1. 기초공사", "0102. 기초공사")
SECTION_HEADER_PATTERN = re.compile(r"^\d+([.\-]\d+)*\.?\s*[가-힣]")

# ── P2-13(2026-08-17): 그룹청킹 L1/L2 대상 doc_type ──
# 방법론 문서 4-2절: 표형 3종만 그룹(L1) 태깅 + 행(L2) 역참조로 전환.
# 공사공정표(공종 1행)·기자재현황/품셈/가이드라인(품목 1개)은 v1 단위 유지.
GROUPED_DOC_TYPES = {"공사비내역서", "설계예산서", "수량산출서"}


def looks_like_noise(text):
    return len(NOISE_PATTERN.findall(text)) >= 2


def is_header_row(cells):
    if not cells or len(cells) > 2:
        return False
    first = cells[0].strip()
    return bool(SECTION_HEADER_PATTERN.match(first))


def tag_topics(text):
    return [name for name, pat in TOPIC_PATTERNS.items() if pat.search(text)]


def engine_link_for(doc_type, text):
    links = []
    if doc_type in ("공사비내역서", "설계예산서", "수량산출서") and CAPEX_KEYWORD.search(text):
        links.append("CAPEX_MAJOR_CATEGORIES")
    if doc_type == "구조계산서" and STRUCT_KEYWORD.search(text):
        links += ["SPEC_TABLE", "REGION_DESIGN_LOAD"]
    return links or None


# ── 컨설팅 단계 태깅 (2026-07-25 추가) ──
# 목적: 각 청크가 스마트팜 시설컨설팅 5단계(진단·설계·비용최적화·운영최적화·유지보수)
# 중 어디에 쓰이는지 표시해, 컨설턴트가 단계별로 근거 청크를 바로 필터링하게 한다.
_C = {
    "진단": re.compile(r"검토|현황|진단|적정성|안전성|적합성|하중|적설|풍하중|구조계산|노후|애로|문제점"),
    "설계": re.compile(r"설계|규격|사양|기준|도면|배치|평면|골조|단면|상세도|시방|KDS|구조검토|용량|제원"),
    "비용최적화": re.compile(r"공사비|내역|단가|금액|수량|재료비|노무비|경비|원가|견적|사업비|보조|자부담|VAT|합계"),
    "운영최적화": re.compile(r"운영|운전|가동|에너지|전기요금|연료비|난방비|인건비|운영비|생산성|수확량|상품률|가동률|자동화|ICT|제어"),
    "유지보수": re.compile(r"하자|보수|사후관리|유지관리|점검|보증|A\s*/?\s*S|내구연한|소모품|교체|정기점검|워런티"),
}
# doc_type만으로도 걸리는 기본 단계(내용 키워드가 약해도 문서 성격상 해당되는 것)
_DOCTYPE_STAGE = {
    "공사비내역서": ["비용최적화", "설계"],
    "설계예산서": ["비용최적화", "설계"],
    "수량산출서": ["비용최적화", "설계"],
    "공사공정표": ["비용최적화"],
    "설계도면": ["설계", "진단"],
    "시방서": ["설계", "진단"],
    "구조계산서": ["진단", "설계"],
    "검토서": ["진단"],
    "기자재현황/품셈/가이드라인": ["비용최적화", "설계"],
}


def consulting_tags(doc_type, text):
    tags = set(_DOCTYPE_STAGE.get(doc_type, []))
    for stage, pat in _C.items():
        if pat.search(text):
            tags.add(stage)
    # 단계 우선순위 순으로 정렬해 결정적 출력
    order = ["진단", "설계", "비용최적화", "운영최적화", "유지보수"]
    return [s for s in order if s in tags]


# ── 성능사양 수치 추출 (견적용) ──
# 수치+단위 토큰을 뽑는다. 단위 뒤에 '원'이 오는 금액은 성능이 아니므로 제외.
SPEC_UNIT = re.compile(
    r"(\d[\d,]*\.?\d*)\s?"
    r"(kcal/?h|kcal|kW/?h|kW|W/㎡·?K|W/m2·?K|W/㎡K|W|kV·?A|kVA|V|A|Hz|"
    r"㎥/?h|m3/?h|㎥/?분|㎥|m3|LPM|L/?min|㎡|m2|mm|cm|Φ|파이|"
    r"℃|°C|kPa|Pa|㎩|%|ppm|μmol|umol|lx|lux|dB|HP|마력|rpm|톤|ton|kg|평)",
    re.I,
)


# ── 비용 수치 추출 (견적/비용최적화용) ──
# 천단위 콤마가 있는 금액/단가 후보. 단, 숫자 바로 뒤에 성능단위가 붙은 값
# (예: 6,000㎥/h, 3,971㎡)은 성능사양이지 금액이 아니므로 제외한다.
MONEY = re.compile(r"\d{1,3}(?:,\d{3})+")


def extract_spec_cost(text):
    specs, spec_starts = [], set()
    for m in SPEC_UNIT.finditer(text):
        spec_starts.add(m.start(1))
        tok = (m.group(1) + m.group(2)).strip()
        if tok not in specs:
            specs.append(tok)
        if len(specs) >= 12:
            break
    costs = []
    for m in MONEY.finditer(text):
        if m.start() in spec_starts:      # 성능수치(6,000㎥ 등)는 금액에서 제외
            continue
        v = m.group(0)
        if v not in costs:
            costs.append(v)
        if len(costs) >= 12:
            break
    return specs, costs


def clean(t):
    return re.sub(r"\s+", " ", t or "").strip()


class ChunkWriter:
    def __init__(self):
        self.chunks = []
        self.skipped_files = []  # (path, reason)

    def add(self, case_name, doc_type, doc_subtype, source_file, page_or_section,
            text, evidence_status="실측", extraction_quality="normal", duplicate_of=None,
            section_context=None, chunk_level=None, parent_group_id=None,
            group_member_rows=None, suppress_topic_tags=False):
        text = clean(text)
        noise = looks_like_noise(text)
        tagging_text = text if not noise else ""
        if section_context:
            tagging_text = clean(section_context) + " " + tagging_text
        chunk_id = f"{case_name}_{doc_type}_{doc_subtype or ''}_{page_or_section}".replace(" ", "")
        rec = {
            "chunk_id": chunk_id,
            "source_file": os.path.relpath(source_file, ROOT).replace("\\", "/"),
            "case_name": case_name,
            "doc_type": doc_type,
            "doc_subtype": doc_subtype,
            "page_or_section": page_or_section,
            "section_context": clean(section_context) if section_context else None,
            # P2-13: L2 행은 주제축 태깅을 시도하지 않는다(태깅은 소속 L1 그룹에서).
            "topic_tags": [] if suppress_topic_tags else tag_topics(tagging_text),
            "consulting_tags": consulting_tags(doc_type, tagging_text),
            "content_summary": text[:200],
            "text_len": len(text),
            "evidence_status": evidence_status,
            "extraction_quality": "noise_suspected" if noise else extraction_quality,
            "engine_link": engine_link_for(doc_type, tagging_text),
            "duplicate_of": duplicate_of,
        }
        # P2-13 그룹청킹 필드 — 해당될 때만 기록(비표형 청크의 스키마는 불변)
        if chunk_level is not None:
            rec["chunk_level"] = chunk_level
        if parent_group_id is not None:
            rec["parent_group_id"] = parent_group_id
        if group_member_rows is not None:
            rec["group_member_rows"] = group_member_rows
        specs, costs = extract_spec_cost(text)
        rec["spec_values"] = specs
        rec["cost_values"] = costs
        rec["spec_signal"] = bool(specs)
        rec["cost_signal"] = bool(costs)
        self.chunks.append(rec)
        return rec

    def skip(self, path, reason):
        self.skipped_files.append((os.path.relpath(path, ROOT).replace("\\", "/"), reason))


class GroupEmitter:
    """P2-13(2026-08-17) — 표형 문서(GROUPED_DOC_TYPES)의 L1 그룹/L2 행 방출기.

    방법론 문서 4-1절: 태깅은 L1(공종/절 묶음)에서 하고, 개별 내역 행(L2)은
    parent_group_id로 그룹에 매달아 두되 주제축 태깅을 시도하지 않는다(부품명
    단일행 미분류의 주원인 제거). 그룹 경계는 is_header_row()가 여는 공종 헤더.
    헤더 없이 시작하는 연속 행 구간은 "그룹(코드나열)" 1그룹으로 묶는다 —
    태깅은 시도하되 대부분 미분류로 남는 게 정상(억지 태깅 금지).

    사용: 헤더 행이면 header_row(), 내역 행이면 row(). 시트/파일 경계에서
    flush() — 버퍼된 그룹을 L1(group_rollup) 1건 + L2 행들로 방출한다.
    L1이 먼저 방출되고 그 chunk_id가 L2의 parent_group_id로 들어간다.
    """

    def __init__(self, w, case_name, doc_type, source_path,
                 evidence_status="실측", duplicate_of=None):
        self.w = w
        self.case_name = case_name
        self.doc_type = doc_type
        self.source_path = source_path
        self.evidence_status = evidence_status
        self.duplicate_of = duplicate_of
        self.n_groups = 0
        self._header = None          # (row_id, text) — 현재 그룹의 공종 헤더
        self._rows = []              # [(row_id, text, quality)]

    def header_row(self, row_id, text):
        self.flush()
        self._header = (row_id, clean(text))

    def row(self, row_id, text, quality):
        self._rows.append((row_id, clean(text), quality))

    def flush(self):
        if self._header is None and not self._rows:
            return
        self.n_groups += 1
        header_text = self._header[1] if self._header else None
        anchor = self._header[0] if self._header else self._rows[0][0]
        pos = f"{anchor}-g{self.n_groups}"
        subtype = "그룹" if header_text else "그룹(코드나열)"
        member_ids = [rid for rid, _, _ in self._rows]
        group_text = " | ".join(([header_text] if header_text else [])
                                + [t for _, t, _ in self._rows])
        grec = self.w.add(
            self.case_name, self.doc_type, subtype, self.source_path, pos,
            group_text, evidence_status=self.evidence_status,
            extraction_quality="group_rollup", duplicate_of=self.duplicate_of,
            section_context=header_text, chunk_level="group",
            group_member_rows=member_ids)
        for rid, text, quality in self._rows:
            self.w.add(
                self.case_name, self.doc_type, None, self.source_path, rid,
                text, evidence_status=self.evidence_status,
                extraction_quality=quality, duplicate_of=self.duplicate_of,
                section_context=header_text, chunk_level="row",
                parent_group_id=grec["chunk_id"], suppress_topic_tags=True)
        self._header = None
        self._rows = []


def chunk_pdf_by_page(w, path, case_name, doc_type, doc_subtype, evidence_status="실측",
                       quality_note=None, duplicate_of=None):
    # 페이지 단위(도면·시방서·구조계산 등)는 pypdf로 텍스트만 빠르게 추출한다.
    # pdfplumber는 벡터가 많은 도면 PDF에서 페이지당 수십 초~수 분이 걸려 비현실적이다.
    # pypdf 추출이 통째로 비면(스캔 등) pdfplumber로 폴백한다.
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        pages = reader.pages
        n_pages = len(pages)
        n_blank = 0
        got_any = False
        for i, page in enumerate(pages):
            try:
                t = page.extract_text() or ""
            except Exception:
                t = ""
            if not t.strip():
                n_blank += 1
                continue
            got_any = True
            w.add(case_name, doc_type, doc_subtype, path, f"p{i+1}", t,
                  evidence_status=evidence_status,
                  extraction_quality=quality_note or "normal",
                  duplicate_of=duplicate_of)
        if got_any or n_pages == 0:
            return n_pages, n_blank
    except Exception:
        pass
    # 폴백: pdfplumber
    with pdfplumber.open(path) as pdf:
        n_pages = len(pdf.pages)
        n_blank = 0
        for i, page in enumerate(pdf.pages):
            t = page.extract_text() or ""
            if not t.strip():
                n_blank += 1
                continue
            w.add(case_name, doc_type, doc_subtype, path, f"p{i+1}", t,
                  evidence_status=evidence_status,
                  extraction_quality=quality_note or "normal",
                  duplicate_of=duplicate_of)
    return n_pages, n_blank


def chunk_pdf_table_or_lines(w, path, case_name, doc_type, doc_subtype, evidence_status="실측",
                              duplicate_of=None):
    n_pages_with_table = 0
    n_pages_line_fallback = 0
    current_section = None
    # P2-13: 표형 3종은 그룹(L1)+행(L2)로 방출. 헤더 캐리포워드가 파일 전체에
    # 걸치는 기존 의미를 따라 emitter도 파일당 1개(페이지·표 경계를 넘어 지속).
    ge = GroupEmitter(w, case_name, doc_type, path, evidence_status, duplicate_of) \
        if doc_type in GROUPED_DOC_TYPES else None
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            if tables:
                n_pages_with_table += 1
                for ti, table in enumerate(tables):
                    for ri, row in enumerate(table):
                        cells = [clean(c) for c in row if c]
                        if not cells:
                            continue
                        row_text = " | ".join(cells)
                        if len(row_text) < 2:
                            continue
                        rid = f"p{i+1}-t{ti+1}-r{ri+1}"
                        if ge is not None:
                            if is_header_row(cells):
                                ge.header_row(rid, row_text)
                            else:
                                ge.row(rid, row_text, "table_extract")
                            continue
                        if is_header_row(cells):
                            current_section = row_text
                            ctx = None
                        else:
                            ctx = current_section
                        w.add(case_name, doc_type, doc_subtype, path, rid,
                              row_text, evidence_status=evidence_status,
                              extraction_quality="table_extract", duplicate_of=duplicate_of,
                              section_context=ctx)
            else:
                t = page.extract_text() or ""
                if not t.strip():
                    continue
                n_pages_line_fallback += 1
                for li, line in enumerate(t.split("\n")):
                    if len(line.strip()) < 2:
                        continue
                    rid = f"p{i+1}-line{li+1}"
                    if ge is not None:
                        if is_header_row([line.strip()]):
                            ge.header_row(rid, line.strip())
                        else:
                            ge.row(rid, line, "line_fallback")
                        continue
                    if is_header_row([line.strip()]):
                        current_section = line.strip()
                        ctx = None
                    else:
                        ctx = current_section
                    w.add(case_name, doc_type, doc_subtype, path, rid,
                          line, evidence_status=evidence_status,
                          extraction_quality="line_fallback", duplicate_of=duplicate_of,
                          section_context=ctx)
    if ge is not None:
        ge.flush()
    return n_pages_with_table, n_pages_line_fallback


def chunk_xlsx(w, path, case_name, doc_type, doc_subtype, evidence_status="실측"):
    wb = openpyxl.load_workbook(path, data_only=True)
    n_rows = 0
    grouped = doc_type in GROUPED_DOC_TYPES  # P2-13
    for ws in wb.worksheets:
        current_section = None
        # 기존 캐리포워드가 시트마다 리셋되므로 emitter도 시트당 1개
        ge = GroupEmitter(w, case_name, doc_type, path, evidence_status) if grouped else None
        for ri, row in enumerate(ws.iter_rows(values_only=True), start=1):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if not cells:
                continue
            row_text = " | ".join(cells)
            if len(row_text) < 2:
                continue
            n_rows += 1
            rid = f"{ws.title}-r{ri}"
            if ge is not None:
                if is_header_row(cells):
                    ge.header_row(rid, row_text)
                else:
                    ge.row(rid, row_text, "xlsx_row")
                continue
            if is_header_row(cells):
                current_section = row_text
                ctx = None
            else:
                ctx = current_section
            w.add(case_name, doc_type, doc_subtype, path, rid,
                  row_text, evidence_status=evidence_status, extraction_quality="xlsx_row",
                  section_context=ctx)
        if ge is not None:
            ge.flush()
    return n_rows


def _xls_sheets_lenient(path):
    """P2-16(2026-08-17) — xlrd 정상 경로가 실패하는 손상 .xls 판독 우회.

    실측 진단(수현건설·최선동·한수진·우민재 공정표·ICT확산내역서 5건): 전부
    진짜 OLE2 바이너리이며, 실패 원인은 셀 데이터가 아니라 메타데이터 손상
    2종이었다 — ①SUPBOOK(외부참조) 문자열의 utf-16 위반, ②NAME(정의된 이름)
    수식의 순환 참조("Excessive indirect references"). 따라서 ①문자열 디코드를
    errors='replace'로 완화하고 ②이름 수식 평가를 건너뛰면 셀 데이터는 온전히
    읽힌다. 패치는 이 함수 안에서만 적용하고 finally로 원상복구한다(정상 파일의
    strict 파싱 경로에 영향 없음). 반환: {시트명: [[셀,...],...]} 또는 None.
    """
    import xlrd
    import xlrd.biffh
    import xlrd.book
    import xlrd.formula
    import xlrd.timemachine

    lenient = lambda b, enc: b.decode(enc, errors="replace")
    noop = lambda *a, **k: None
    saved = [(m, "unicode", m.unicode) for m in (xlrd.timemachine, xlrd.biffh, xlrd.book)]
    saved += [(xlrd.formula, "evaluate_name_formula", xlrd.formula.evaluate_name_formula),
              (xlrd.book, "evaluate_name_formula", xlrd.book.evaluate_name_formula)]
    try:
        for m in (xlrd.timemachine, xlrd.biffh, xlrd.book):
            m.unicode = lenient
        xlrd.formula.evaluate_name_formula = noop
        xlrd.book.evaluate_name_formula = noop
        book = xlrd.open_workbook(path)
        out = {}
        for sheet in book.sheets():
            rows = []
            for ri in range(sheet.nrows):
                vals = []
                for c in sheet.row_values(ri):
                    if isinstance(c, float) and c == int(c):
                        c = int(c)
                    vals.append(c)
                rows.append(vals)
            out[sheet.name] = rows
        return out
    except Exception:
        return None
    finally:
        for mod, attr, orig in saved:
            setattr(mod, attr, orig)


def chunk_xls_legacy(w, path, case_name, doc_type, doc_subtype, evidence_status="실측"):
    import pandas as pd
    n_rows = 0
    row_quality = "xls_row"
    try:
        dfs = pd.read_excel(path, sheet_name=None, header=None, engine="xlrd")
        sheets = {name: df.values.tolist() for name, df in dfs.items()}
    except Exception as e:
        # P2-16: strict 실패 시 관대 판독 우회(메타데이터 손상 .xls). 우회로 읽은
        # 행은 품질 라벨을 구분해 정직하게 표시한다.
        sheets = _xls_sheets_lenient(path)
        if sheets is None:
            w.skip(path, f"xls 읽기 실패(관대 판독 우회 포함): {e}")
            return 0
        row_quality = "xls_row_lenient"
    grouped = doc_type in GROUPED_DOC_TYPES  # P2-13
    for sheet_name, rows in sheets.items():
        current_section = None
        ge = GroupEmitter(w, case_name, doc_type, path, evidence_status) if grouped else None
        for ri, row in enumerate(rows):
            cells = [str(c) for c in row if str(c) not in ("nan", "None", "")]
            if not cells:
                continue
            row_text = " | ".join(cells)
            if len(row_text) < 2:
                continue
            n_rows += 1
            rid = f"{sheet_name}-r{ri+1}"
            if ge is not None:
                if is_header_row(cells):
                    ge.header_row(rid, row_text)
                else:
                    ge.row(rid, row_text, row_quality)
                continue
            if is_header_row(cells):
                current_section = row_text
                ctx = None
            else:
                ctx = current_section
            w.add(case_name, doc_type, doc_subtype, path, rid,
                  row_text, evidence_status=evidence_status, extraction_quality=row_quality,
                  section_context=ctx)
        if ge is not None:
            ge.flush()
    return n_rows


def chunk_docx(w, path, case_name, doc_type, doc_subtype, evidence_status="실측"):
    import docx
    d = docx.Document(path)
    n = 0
    current_section = None
    HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Title"}
    for pi, para in enumerate(d.paragraphs):
        t = clean(para.text)
        if len(t) < 2:
            continue
        is_heading = para.style.name in HEADING_STYLES or bool(SECTION_HEADER_PATTERN.match(t))
        if is_heading:
            current_section = t
            ctx = None
        else:
            ctx = current_section
        n += 1
        w.add(case_name, doc_type, doc_subtype, path, f"para{pi+1}", t,
              evidence_status=evidence_status, extraction_quality="docx_para",
              section_context=ctx)
    for ti, table in enumerate(d.tables):
        for ri, row in enumerate(table.rows):
            cells = [clean(c.text) for c in row.cells if clean(c.text)]
            if not cells:
                continue
            row_text = " | ".join(cells)
            n += 1
            w.add(case_name, doc_type, doc_subtype, path, f"t{ti+1}-r{ri+1}",
                  row_text, evidence_status=evidence_status, extraction_quality="docx_table_row")
    return n


def write_outputs(w, jsonl_path, summary_path, extra_log=None):
    with open(jsonl_path, "w", encoding="utf-8") as f:
        for c in w.chunks:
            f.write(json.dumps(c, ensure_ascii=False) + "\n")

    from collections import Counter
    by_doc_type = Counter(c["doc_type"] for c in w.chunks)
    by_case = Counter(c["case_name"] for c in w.chunks)
    by_quality = Counter(c["extraction_quality"] for c in w.chunks)
    topic_counter = Counter()
    for c in w.chunks:
        for t in c["topic_tags"]:
            topic_counter[t] += 1
    untagged = sum(1 for c in w.chunks if not c["topic_tags"])
    linked = sum(1 for c in w.chunks if c["engine_link"])
    # P2-13 태깅 단위 지표 — L2 행(chunk_level=="row")은 설계상 무태깅이므로
    # 미분류율은 "태깅 단위"(행 제외 전체) 기준으로 따로 계산한다.
    n_l2_rows = sum(1 for c in w.chunks if c.get("chunk_level") == "row")
    n_groups = sum(1 for c in w.chunks if c.get("chunk_level") == "group")
    units = [c for c in w.chunks if c.get("chunk_level") != "row"]
    untagged_units = sum(1 for c in units if not c["topic_tags"])
    g_units = [c for c in units if c["doc_type"] in GROUPED_DOC_TYPES]
    g_untagged = sum(1 for c in g_units if not c["topic_tags"])
    consult_counter = Counter()
    for c in w.chunks:
        for t in c.get("consulting_tags", []):
            consult_counter[t] += 1
    spec_n = sum(1 for c in w.chunks if c.get("spec_signal"))
    cost_n = sum(1 for c in w.chunks if c.get("cost_signal"))

    with open(summary_path, "w", encoding="utf-8") as f:
        if extra_log:
            f.write("=== 처리 로그 ===\n")
            for line in extra_log:
                f.write(line + "\n")
            f.write("\n")

        f.write(f"=== 총 청크 수: {len(w.chunks)} ===\n")
        f.write(f"=== 처리 실패/건너뛴 파일 수: {len(w.skipped_files)} ===\n")
        if n_l2_rows or n_groups:
            f.write(f"\n-- P2-13 그룹청킹 지표 --\n")
            f.write(f"L1 그룹: {n_groups} / L2 행(무태깅): {n_l2_rows}\n")
            f.write(f"태깅 단위(행 제외): {len(units)}, 미분류 {untagged_units} "
                    f"({untagged_units/max(len(units),1)*100:.1f}%)\n")
            f.write(f"표형 3종 태깅 단위: {len(g_units)}, 미분류 {g_untagged} "
                    f"({g_untagged/max(len(g_units),1)*100:.1f}%) — 목표 30% 이하\n")

        f.write("\n-- doc_type별 청크 수 --\n")
        for k, v in by_doc_type.most_common():
            f.write(f"  {k}: {v}\n")
        f.write("\n-- case별 청크 수 (상위 30) --\n")
        for k, v in by_case.most_common(30):
            f.write(f"  {k}: {v}\n")
        f.write("\n-- extraction_quality별 청크 수 --\n")
        for k, v in by_quality.most_common():
            f.write(f"  {k}: {v}\n")
        f.write("\n-- 주제축별 태깅 건수 (청크 1개가 복수 축에 걸릴 수 있음) --\n")
        for k, v in topic_counter.most_common():
            f.write(f"  {k}: {v}\n")
        f.write(f"\n미분류(태그 0개) 청크 수: {untagged} / {len(w.chunks)}"
                f" ({untagged/len(w.chunks)*100:.1f}%)\n" if w.chunks else "\n청크 없음\n")
        f.write(f"\nengine_link 있음: {linked} / {len(w.chunks)}\n")

        f.write("\n-- 컨설팅단계별 태깅 건수 (복수 단계 가능) --\n")
        for k in ["진단", "설계", "비용최적화", "운영최적화", "유지보수"]:
            f.write(f"  {k}: {consult_counter.get(k, 0)}\n")
        f.write(f"\n성능사양 수치 포함(spec_signal) 청크: {spec_n} / {len(w.chunks)}\n")
        f.write(f"비용 수치 포함(cost_signal) 청크: {cost_n} / {len(w.chunks)}\n")

        if w.skipped_files:
            f.write("\n-- 건너뛴 파일 (근거 없이 채우지 않고 명시적으로 제외) --\n")
            for path, reason in w.skipped_files:
                f.write(f"  {path}: {reason}\n")

    print(f"완료: 청크 {len(w.chunks)}개 -> {jsonl_path}")
    print(f"건너뛴 파일 {len(w.skipped_files)}개")
    print(f"요약 -> {summary_path}")
